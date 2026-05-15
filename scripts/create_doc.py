from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_article_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('介護現場を救うAI評価システム『ルミタス』を作ってみた。現場の「頑張り」を可視化するための挑戦', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('「あぁ、また今日も記録が終わらない……」')
    doc.add_paragraph('介護や福祉の現場で、そんなため息をついたことはありませんか？')
    doc.add_paragraph('利用者の皆さんと笑顔で接している時間と同じくらい、あるいはそれ以上に、私たちは「書類」と向き合っています。')
    doc.add_paragraph('そんな現場の景色を変えたくて、今回、私は新しい職員評価システム『ルミタス（Lumitas）』をゼロから作り上げました。今回は、その開発の裏側にある「こだわり」と、ルミタスが目指す「新しい評価の形」についてお話しします。')

    doc.add_heading('1. 「ミスを減らす」から「発見を褒める」へ', level=1)
    doc.add_paragraph('これまでの評価システムは、どちらかというと「何ができなかったか」を見つけ出すものになりがちでした。特にインシデント（ヒヤリハット）報告などは、書く側も「怒られるのではないか」と身構えてしまうものです。')
    doc.add_paragraph('ルミタスは、ここを180度変えました。')
    doc.add_paragraph('「ヒヤリハットを見つけたこと」そのものを、AIがプラスの評価としてカウントします。')
    
    # Image Marker 1
    p = doc.add_paragraph()
    run = p.add_run('【ここに画像：lumitas_point1_good_catch_v6_safe_support を挿入】')
    run.bold = True
    run.font.color.rgb = None # Standard black

    doc.add_paragraph('重大な事故を未然に防ぐための「気づき」は、施設にとって最大の財産。報告すればするほど、その職員の「安全への意識」がスコアとして加算される仕組みにしました。ミスを隠す文化ではなく、みんなで共有して安全を高める文化を、システムから作りたかったのです。')

    doc.add_heading('2. 「ありがとう」がポイントになる。助け合いの可視化', level=1)
    doc.add_paragraph('介護はチームプレーです。でも、同僚をフォローしたり、ちょっとした手伝いをした時の頑張りは、これまでは評価シートに載りにくいものでした。')
    doc.add_paragraph('そこでルミタスには、「サンクスバッジ（ピアボーナス）」という機能を搭載しました。')
    
    # Image Marker 2
    p = doc.add_paragraph()
    run = p.add_run('【ここに画像：lumitas_point2_thanks_badge_v3 を挿入】')
    run.bold = True

    doc.add_paragraph('スタッフ同士で「さっきの介助、助かったよ！」とバッジを送ることができ、この「感謝の数」も評価に反映されます。「誰かが見ていてくれる」という安心感が、現場の助け合いを加速させます。')

    doc.add_heading('3. 「書く」ストレスを、AIが解消する', level=1)
    doc.add_paragraph('忙しい現場で「入力」に時間を取られるのは本末転倒です。')
    doc.add_paragraph('そこで導入したのが、「AIによる音声記録」。')
    
    # Image Marker 3
    p = doc.add_paragraph()
    run = p.add_run('【ここに画像：lumitas_point3_ai_voice を挿入】')
    run.bold = True

    doc.add_paragraph('スマホに向かってつぶやくだけで、AIがその言葉を整え、人事考課にも使える立派な記録へと自動変換してくれます。キーボードを叩く時間を減らし、その分を利用者さんと向き合う時間に使ってほしい。そんな想いを込めています。')

    doc.add_heading('4. 「鉄壁の守り」と「直感的な操作」の両立', level=1)
    doc.add_paragraph('個人情報を扱うため、セキュリティにも妥協しませんでした。')
    
    # Image Marker 4
    p = doc.add_paragraph()
    run = p.add_run('【ここに画像：lumitas_point4_security_ui_v2 を挿入】')
    run.bold = True

    doc.add_paragraph('Googleのシステムを賢く活用した「独自の2段階認証」を開発し、コストを抑えつつ銀行並みの安心感を実現しています。また、スマホでシュッと「スワイプ」するだけで目標更新ができる直感的な操作感も、フルスクラッチ開発だからこそ実現できたこだわりです。')

    doc.add_heading('おわりに：ルミタスが見つめる、その先', level=1)
    doc.add_paragraph('ルミタス（Lumitas）という名前には、現場の「光（Lumi）」を「足す（+）」という意味を込めています。')
    doc.add_paragraph('これまで誰にも気づかれなかったかもしれない「小さな光」を、AIの力でしっかりと拾い上げたい。「頑張っている人が、正しく報われる現場」を目指して、ルミタスの挑戦は続きます。')
    doc.add_paragraph('次回の記事では、具体的な画面などもお見せしながら、詳しい機能についてご紹介します！')

    doc.save('lumitas_note_article.docx')
    print('lumitas_note_article.docx created successfully.')

if __name__ == '__main__':
    create_article_doc()
